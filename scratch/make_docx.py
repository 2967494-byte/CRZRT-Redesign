import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()

# Page margins: 2 cm all around
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)

# Base normal style font: Arial, 11 pt
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
font.color.rgb = RGBColor(0x22, 0x22, 0x22)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run('КРАТКОЕ ОПИСАНИЕ ИНФОРМАЦИОННОЙ СИСТЕМЫ\n')
run_title.bold = True
run_title.font.size = Pt(14)
run_title.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
p_title.paragraph_format.space_after = Pt(12)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 11)
    run.font.color.rgb = RGBColor(0x0F, 0x5B, 0x30)
    return p

# 1. Общие сведения
add_heading('1. Общие сведения об информационной системе и Операторе')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.add_run('• Полное наименование ИС: ').bold = True
p.add_run('Автоматизированная информационная система «Модуль Assessment — профессиональное тестирование специалистов сферы закупок».\n')
p.add_run('• Краткое наименование ИС: ').bold = True
p.add_run('ИС «Модуль Assessment».\n')
p.add_run('• Оператор информационной системы: ').bold = True
p.add_run('Акционерное общество «Центр развития закупок Республики Татарстан» (АО «ЦРЗ РТ»).\n')
p.add_run('• ИНН / КПП Оператора: ').bold = True
p.add_run('1655291244 / 165501001\n')
p.add_run('• ОГРН Оператора: ').bold = True
p.add_run('1141690026217\n')
p.add_run('• Юридический / почтовый адрес: ').bold = True
p.add_run('420107, Российская Федерация, Республика Татарстан, г. Казань, ул. Петербургская, д. 86.\n')
p.add_run('• Сетевой адрес (URL) ИС: ').bold = True
p.add_run('https://test.zakupki.tatar (основной портал: https://zakupki.tatar)\n')
p.add_run('• Электронная почта поддержки: ').bold = True
p.add_run('crz.rt@tatar.ru')

# 2. Назначение и правовые основания
add_heading('2. Назначение и правовые основания функционирования ИС')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.add_run('Назначение информационной системы: ').bold = True
p.add_run('Информационная система предназначена для автоматизации процессов независимой квалификационной оценки, профессионального тестирования и мониторинга уровня компетенций специалистов в сфере государственных, муниципальных и корпоративных закупок (включая контрактных управляющих, сотрудников контрактных служб, членов закупочных комиссий и представителей организаций-заказчиков Республики Татарстан и других субъектов РФ).')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.add_run('Правовые основания создания и функционирования:\n').bold = True
p.add_run('1. Федеральный закон от 05.04.2013 № 44-ФЗ «О контрактной системе в сфере закупок товаров, работ, услуг для обеспечения государственных и муниципальных нужд» (ст. 9 — принцип профессионализма заказчика);\n')
p.add_run('2. Федеральный закон от 27.07.2006 № 149-ФЗ «Об информации, информационных технологиях и о защите информации»;\n')
p.add_run('3. Федеральный закон от 27.07.2006 № 152-ФЗ «О персональных данных»;\n')
p.add_run('4. Постановление Правительства РФ от 28.11.2011 № 977 «О единой системе идентификации и аутентификации в инфраструктуре, обеспечивающей информационно-технологическое взаимодействие информационных систем, используемых для предоставления государственных и муниципальных услуг в электронной форме».')

# 3. Основные функции
add_heading('3. Основные функции информационной системы')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.add_run('ИС «Модуль Assessment» обеспечивает выполнение следующего комплекса задач:\n')
p.add_run('1. Идентификация и аутентификация пользователей: ').bold = True
p.add_run('обеспечение санкционированного и защищенного входа в личный кабинет, в том числе с использованием ЕСИА;\n')
p.add_run('2. Ведение профиля и учет ведомственной принадлежности: ').bold = True
p.add_run('хранение данных об участнике тестирования, подтверждение связи сотрудника с организацией-работодателем (заказчиком);\n')
p.add_run('3. Организация и проведение аттестационных кампаний: ').bold = True
p.add_run('предоставление доступа к назначенным квалификационным тестам в строгих временных рамках регламентных периодов;\n')
p.add_run('4. Контроль тестирования в электронной форме: ').bold = True
p.add_run('динамическая генерация индивидуальных тестовых билетов из верифицированного банка вопросов, тайминг прохождения, проверка правильности ответов на серверной стороне;\n')
p.add_run('5. Обработка и фиксация результатов: ').bold = True
p.add_run('автоматизированный подсчет баллов, определение статуса сдачи теста, протоколирование попыток и исключение повторного несанкционированного прохождения;\n')
p.add_run('6. Аналитическая отчетность: ').bold = True
p.add_run('формирование ведомственной сводной статистики по организациям, отраслям и муниципальным образованиям, выявление сложных тем для корректировки программ повышения квалификации.')

# 4. Категории пользователей
add_heading('4. Категории пользователей информационной системы')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.add_run('• Физические лица (участники): ').bold = True
p.add_run('контрактные управляющие, руководители и специалисты контрактных служб, члены закупочных комиссий государственных и муниципальных заказчиков, слушатели образовательных программ.\n')
p.add_run('• Администраторы и модераторы ИС: ').bold = True
p.add_run('уполномоченные сотрудники АО «Центр развития закупок Республики Татарстан», осуществляющие методическое сопровождение тестирования и контроль результатов.')

# 5. Интеграция с ЕСИА
add_heading('5. Интеграция с ЕСИА и состав запрашиваемых сведений')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.add_run('Интеграция осуществляется по протоколу OpenID Connect / OAuth 2.0 в соответствии с Регламентом информационно-технологического взаимодействия с ЕСИА.\n\n')
p.add_run('В рамках соблюдения принципа минимизации и соразмерности обработки персональных данных (ч. 2 ст. 5 Федерального закона № 152-ФЗ), информационная система запрашивает из ЕСИА исключительно базовый набор сведений, строго необходимый для однозначной идентификации участника и ведения персонального учета результатов аттестации:')

# Таблица прав
table = doc.add_table(rows=5, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

headers = ['Запрашиваемый атрибут / Scope', 'Назначение использования в ИС']
data = [
    ['Системный идентификатор пользователя (OID / sub)', 'Однозначная связка учетной записи ЕСИА с профилем в ИС «Модуль Assessment»'],
    ['Фамилия, имя, отчество (fullname / openid)', 'Идентификация участника в протоколах тестирования и сертификатах/удостоверениях'],
    ['Адрес электронной почты (email)', 'Направление уведомлений о назначенных тестированиях и результатах'],
    ['Номер мобильного телефона (mobile)', 'Оперативная связь и дополнительный канал подтверждения доступа']
]

for col_idx, text in enumerate(headers):
    cell = table.cell(0, col_idx)
    cell.text = text
    shading = parse_xml(r'<w:shd {} w:fill="F0F4F8"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading)
    p = cell.paragraphs[0]
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(10)

for row_idx, row_data in enumerate(data, start=1):
    for col_idx, text in enumerate(row_data):
        cell = table.cell(row_idx, col_idx)
        cell.text = text
        p = cell.paragraphs[0]
        p.runs[0].font.size = Pt(10)

# Set col widths
for row in table.rows:
    row.cells[0].width = Inches(2.8)
    row.cells[1].width = Inches(3.9)

# 6. Цели интеграции
add_heading('6. Цели интеграции с ЕСИА')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.add_run('1. Достоверность идентификации: ').bold = True
p.add_run('исключение прохождения квалификационных тестов поддельными, анонимными или неуполномоченными лицами;\n')
p.add_run('2. Единая точка входа (SSO): ').bold = True
p.add_run('обеспечение удобного и безопасного бесшовного доступа пользователей без необходимости администрирования отдельных небезопасных паролей;\n')
p.add_run('3. Персональная ответственность и юридическая значимость: ').bold = True
p.add_run('обеспечение легитимности формируемых результатов профессиональной оценки специалистов контрактной системы;\n')
p.add_run('4. Информационная безопасность: ').bold = True
p.add_run('снижение рисков компрометации учетных записей за счет применения стандартов защиты и двухфакторной аутентификации ЕСИА.')

# 7. Среда интеграции
add_heading('7. Среда интеграции')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
p.add_run('На текущем этапе интеграция осуществляется в ')
p.add_run('тестовой среде ЕСИА ').bold = True
p.add_run('(esia-portal1.test.gosuslugi.ru) с целью выполнения технической отладки, конфигурирования клиентского приложения и сквозной проверки сценариев аутентификации перед выводом функционала в промышленную эксплуатацию.')

# Sign block
p_sign = doc.add_paragraph()
p_sign.paragraph_format.keep_with_next = True
p_sign.add_run('Оператор ИС:\n').bold = True
p_sign.add_run('АО «Центр развития закупок Республики Татарстан»\n\n')
p_sign.add_run('Генеральный директор / Уполномоченное лицо:\n')
p_sign.add_run('____________________ / ____________________ /\n\n')
p_sign.add_run('М.П.\n\n')
p_sign.add_run('«___» ____________ 2026 г.')

doc.save(r'Документы\Краткое_описание_ИС_ЕСИА.docx')
print('Successfully saved DOCX!')
